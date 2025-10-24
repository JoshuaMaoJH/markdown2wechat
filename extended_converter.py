#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
扩展版Markdown到微信公众号转换器
支持多种输入格式：Markdown, HTML, TXT, RST, AsciiDoc, Word
"""

import markdown
import re
from bs4 import BeautifulSoup
import argparse
import sys
from pathlib import Path
from wechat_styles import WeChatStyleTemplates

# 尝试导入可选依赖
try:
    import docutils.core
    import docutils.writers.html4css1
    RST_AVAILABLE = True
except ImportError:
    RST_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import striprtf
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False


class ExtendedMarkdownToWeChatConverter:
    """扩展版Markdown到微信公众号格式转换器"""
    
    def __init__(self, style="default"):
        """初始化转换器"""
        self.style = style
        self.wechat_styles = WeChatStyleTemplates.get_style_template(style)
        
        # Markdown配置
        self.md_extensions = [
            'markdown.extensions.tables',
            'markdown.extensions.fenced_code',
            'markdown.extensions.codehilite',
            'markdown.extensions.toc',
            'markdown.extensions.extra',
        ]
        
        self.md_config = {
            'markdown.extensions.codehilite': {
                'css_class': 'highlight',
                'use_pygments': True,
                'noclasses': True,
            }
        }
    
    def detect_file_format(self, file_path):
        """检测文件格式"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        format_map = {
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.txt': 'text',
            '.rst': 'rst',
            '.docx': 'docx',
            '.rtf': 'rtf',
        }
        
        return format_map.get(suffix, 'unknown')
    
    def convert_markdown(self, content):
        """转换Markdown格式"""
        md = markdown.Markdown(
            extensions=self.md_extensions,
            extension_configs=self.md_config
        )
        return md.convert(content)
    
    def convert_html(self, content):
        """转换HTML格式"""
        # HTML可以直接使用，只需要优化样式
        soup = BeautifulSoup(content, 'html.parser')
        return str(soup)
    
    def convert_text(self, content):
        """转换纯文本格式"""
        # 将纯文本转换为简单的HTML
        lines = content.split('\n')
        html_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                html_lines.append('<br>')
            else:
                # 简单的格式检测
                if line.startswith('# '):
                    html_lines.append(f'<h1>{line[2:]}</h1>')
                elif line.startswith('## '):
                    html_lines.append(f'<h2>{line[3:]}</h2>')
                elif line.startswith('### '):
                    html_lines.append(f'<h3>{line[4:]}</h3>')
                elif line.startswith('- '):
                    html_lines.append(f'<li>{line[2:]}</li>')
                else:
                    html_lines.append(f'<p>{line}</p>')
        
        return '\n'.join(html_lines)
    
    def convert_rst(self, content):
        """转换RST格式"""
        if not RST_AVAILABLE:
            raise ImportError("需要安装 docutils: pip install docutils")
        
        # 使用docutils转换RST到HTML
        html = docutils.core.publish_parts(
            source=content,
            writer_name='html',
            settings_overrides={'output_encoding': 'unicode'}
        )['html_body']
        
        return html
    
    def convert_docx(self, file_path):
        """转换Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        html_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # 简单的样式检测
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.split()[-1]
                    if level.isdigit():
                        html_content.append(f'<h{level}>{text}</h{level}>')
                    else:
                        html_content.append(f'<h1>{text}</h1>')
                else:
                    html_content.append(f'<p>{text}</p>')
        
        return '\n'.join(html_content)
    
    def convert_rtf(self, content):
        """转换RTF格式"""
        if not RTF_AVAILABLE:
            raise ImportError("需要安装 striprtf: pip install striprtf")
        
        # 使用striprtf转换RTF到纯文本
        text = striprtf.rtf_to_text(content)
        return self.convert_text(text)
    
    def convert_content(self, content, file_format, file_path=None):
        """根据格式转换内容"""
        if file_format == 'markdown':
            return self.convert_markdown(content)
        elif file_format == 'html':
            return self.convert_html(content)
        elif file_format == 'text':
            return self.convert_text(content)
        elif file_format == 'rst':
            return self.convert_rst(content)
        elif file_format == 'docx':
            return self.convert_docx(file_path)
        elif file_format == 'rtf':
            return self.convert_rtf(content)
        else:
            raise ValueError(f"不支持的格式: {file_format}")
    
    def optimize_for_wechat(self, html_content):
        """优化HTML内容以适配微信公众号"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 处理图片标签
        for img in soup.find_all('img'):
            if not img.get('style'):
                img['style'] = 'max-width: 100%; height: auto; display: block; margin: 1em auto;'
        
        # 处理表格
        for table in soup.find_all('table'):
            table['style'] = 'width: 100%; border-collapse: collapse; margin: 1em 0;'
        
        # 处理代码块
        for pre in soup.find_all('pre'):
            pre['style'] = 'background-color: #2c3e50; color: #ecf0f1; padding: 1em; border-radius: 5px; overflow-x: auto;'
        
        # 处理引用块
        for blockquote in soup.find_all('blockquote'):
            blockquote['style'] = 'margin: 1em 0; padding: 0.5em 1em; background-color: #f8f9fa; border-left: 4px solid #3498db;'
        
        return str(soup)
    
    def create_wechat_html(self, html_content, title="", subtitle=""):
        """创建完整的微信公众号HTML文档"""
        optimized_html = self.optimize_for_wechat(html_content)
        
        full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title if title else '微信公众号文章'}</title>
    {self.wechat_styles}
</head>
<body>
"""
        
        if title:
            full_html += f'    <h1 class="wechat-title">{title}</h1>\n'
        if subtitle:
            full_html += f'    <p class="wechat-subtitle">{subtitle}</p>\n'
        
        full_html += f"    {optimized_html}\n"
        full_html += """
</body>
</html>"""
        
        return full_html
    
    def convert_file(self, input_file, output_file=None, title="", subtitle=""):
        """转换文件"""
        try:
            # 检测文件格式
            file_format = self.detect_file_format(input_file)
            
            if file_format == 'unknown':
                print(f"不支持的文件格式: {Path(input_file).suffix}")
                return None
            
            # 读取文件内容
            if file_format == 'docx':
                # Word文档需要特殊处理
                html_content = self.convert_docx(input_file)
            else:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                html_content = self.convert_content(content, file_format, input_file)
            
            # 创建微信公众号HTML
            wechat_html = self.create_wechat_html(html_content, title, subtitle)
            
            # 确定输出文件名
            if not output_file:
                input_path = Path(input_file)
                output_file = input_path.with_suffix('.html')
            
            # 保存HTML文件
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(wechat_html)
            
            print(f"转换完成！")
            print(f"输入文件: {input_file} ({file_format})")
            print(f"输出文件: {output_file}")
            print(f"可以直接复制HTML内容到微信公众号编辑器")
            
            return wechat_html
            
        except Exception as e:
            print(f"转换失败: {str(e)}")
            return None


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='扩展版Markdown到微信公众号文章格式转换器')
    parser.add_argument('input', nargs='?', help='输入文件路径')
    parser.add_argument('-o', '--output', help='输出HTML文件路径')
    parser.add_argument('-t', '--title', help='文章标题')
    parser.add_argument('-s', '--subtitle', help='文章副标题')
    parser.add_argument('--style', help='文章风格', 
                       choices=WeChatStyleTemplates.get_available_styles(),
                       default='default')
    parser.add_argument('--list-styles', action='store_true', help='列出所有可用风格')
    parser.add_argument('--list-formats', action='store_true', help='列出支持的输入格式')
    
    args = parser.parse_args()
    
    # 列出支持的格式
    if args.list_formats:
        print("支持的输入格式：")
        print("=" * 50)
        formats = [
            ("Markdown", ".md, .markdown", "完全支持，包括表格、代码、列表等"),
            ("HTML", ".html, .htm", "直接优化，保持原有结构"),
            ("纯文本", ".txt", "基本格式转换，支持简单标题和列表"),
            ("RST", ".rst", "需要安装 docutils: pip install docutils"),
            ("Word", ".docx", "需要安装 python-docx: pip install python-docx"),
            ("RTF", ".rtf", "需要安装 striprtf: pip install striprtf"),
        ]
        
        for name, extensions, description in formats:
            print(f"{name:10} {extensions:15} - {description}")
        
        print("\n依赖安装命令：")
        print("pip install docutils python-docx striprtf")
        return
    
    # 列出风格
    if args.list_styles:
        print("可用的文章风格：")
        print("=" * 50)
        for style in WeChatStyleTemplates.get_available_styles():
            description = WeChatStyleTemplates.get_style_description(style)
            print(f"{style:12} - {description}")
        return
    
    # 检查输入文件
    if not args.input:
        parser.error("需要提供输入文件路径")
    
    # 创建转换器
    converter = ExtendedMarkdownToWeChatConverter(style=args.style)
    
    # 执行转换
    converter.convert_file(args.input, args.output, args.title, args.subtitle)


if __name__ == "__main__":
    main()
