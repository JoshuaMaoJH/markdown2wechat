#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用格式到微信公众号转换器
采用两步转换策略：其他格式 → Markdown → 微信公众号HTML
"""

import markdown
import re
from bs4 import BeautifulSoup
import argparse
import sys
from pathlib import Path
from wechat_styles import WeChatStyleTemplates

# 尝试导入可选依赖
try:
    import docutils.core
    import docutils.writers.html4css1
    RST_AVAILABLE = True
except ImportError:
    RST_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import striprtf
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    import mammoth
    DOCX_MAMMOTH_AVAILABLE = True
except ImportError:
    DOCX_MAMMOTH_AVAILABLE = False


class UniversalToWeChatConverter:
    """通用格式到微信公众号转换器"""
    
    def __init__(self, style="default"):
        """初始化转换器"""
        self.style = style
        self.wechat_styles = WeChatStyleTemplates.get_style_template(style)
        
        # Markdown配置
        self.md_extensions = [
            'markdown.extensions.tables',
            'markdown.extensions.fenced_code',
            'markdown.extensions.codehilite',
            'markdown.extensions.toc',
            'markdown.extensions.extra',
        ]
        
        self.md_config = {
            'markdown.extensions.codehilite': {
                'css_class': 'highlight',
                'use_pygments': True,
                'noclasses': True,
            }
        }
    
    def detect_file_format(self, file_path):
        """检测文件格式"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        format_map = {
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.txt': 'text',
            '.rst': 'rst',
            '.docx': 'docx',
            '.rtf': 'rtf',
        }
        
        return format_map.get(suffix, 'unknown')
    
    def convert_to_markdown(self, content, file_format, file_path=None):
        """将各种格式转换为Markdown"""
        if file_format == 'markdown':
            return content
        elif file_format == 'html':
            return self.html_to_markdown(content)
        elif file_format == 'text':
            return self.text_to_markdown(content)
        elif file_format == 'rst':
            return self.rst_to_markdown(content)
        elif file_format == 'docx':
            return self.docx_to_markdown(file_path)
        elif file_format == 'rtf':
            return self.rtf_to_markdown(content)
        else:
            raise ValueError(f"不支持的格式: {file_format}")
    
    def html_to_markdown(self, html_content):
        """HTML转Markdown"""
        soup = BeautifulSoup(html_content, 'html.parser')
        markdown_content = []
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'blockquote', 'pre', 'code', 'strong', 'em', 'a', 'img']):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = element.name[1]
                markdown_content.append(f"{'#' * int(level)} {element.get_text()}")
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    markdown_content.append(text)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    markdown_content.append(f"- {li.get_text()}")
            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li', recursive=False), 1):
                    markdown_content.append(f"{i}. {li.get_text()}")
            elif element.name == 'blockquote':
                text = element.get_text().strip()
                markdown_content.append(f"> {text}")
            elif element.name == 'pre':
                code_text = element.get_text()
                markdown_content.append(f"```\n{code_text}\n```")
            elif element.name == 'code':
                text = element.get_text()
                markdown_content.append(f"`{text}`")
            elif element.name == 'strong':
                text = element.get_text()
                markdown_content.append(f"**{text}**")
            elif element.name == 'em':
                text = element.get_text()
                markdown_content.append(f"*{text}*")
            elif element.name == 'a':
                text = element.get_text()
                href = element.get('href', '')
                markdown_content.append(f"[{text}]({href})")
            elif element.name == 'img':
                src = element.get('src', '')
                alt = element.get('alt', '')
                markdown_content.append(f"![{alt}]({src})")
        
        return '\n\n'.join(markdown_content)
    
    def text_to_markdown(self, text_content):
        """纯文本转Markdown"""
        lines = text_content.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append('')
            elif line.startswith('# '):
                markdown_lines.append(f"# {line[2:]}")
            elif line.startswith('## '):
                markdown_lines.append(f"## {line[3:]}")
            elif line.startswith('### '):
                markdown_lines.append(f"### {line[4:]}")
            elif line.startswith('- '):
                markdown_lines.append(f"- {line[2:]}")
            elif line.startswith('* '):
                markdown_lines.append(f"- {line[2:]}")
            elif line.startswith('> '):
                markdown_lines.append(f"> {line[2:]}")
            elif line.startswith('1. '):
                markdown_lines.append(f"1. {line[3:]}")
            else:
                markdown_lines.append(line)
        
        return '\n'.join(markdown_lines)
    
    def rst_to_markdown(self, rst_content):
        """RST转Markdown"""
        if not RST_AVAILABLE:
            raise ImportError("需要安装 docutils: pip install docutils")
        
        # 使用docutils转换RST到HTML，然后转Markdown
        html = docutils.core.publish_parts(
            source=rst_content,
            writer_name='html',
            settings_overrides={'output_encoding': 'unicode'}
        )['html_body']
        
        return self.html_to_markdown(html)
    
    def docx_to_markdown(self, file_path):
        """Word文档转Markdown"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检测标题样式
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.split()[-1]
                if level.isdigit():
                    markdown_content.append(f"{'#' * int(level)} {text}")
                else:
                    markdown_content.append(f"# {text}")
            else:
                # 检测粗体和斜体
                formatted_text = text
                for run in paragraph.runs:
                    if run.bold:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    if run.italic:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
                
                markdown_content.append(formatted_text)
        
        return '\n\n'.join(markdown_content)
    
    def rtf_to_markdown(self, rtf_content):
        """RTF转Markdown"""
        if not RTF_AVAILABLE:
            raise ImportError("需要安装 striprtf: pip install striprtf")
        
        # 使用striprtf转换RTF到纯文本
        text = striprtf.rtf_to_text(rtf_content)
        return self.text_to_markdown(text)
    
    def markdown_to_wechat_html(self, markdown_content, title="", subtitle=""):
        """Markdown转微信公众号HTML"""
        # 转换为HTML
        md = markdown.Markdown(
            extensions=self.md_extensions,
            extension_configs=self.md_config
        )
        html_content = md.convert(markdown_content)
        
        # 优化HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 处理图片标签
        for img in soup.find_all('img'):
            if not img.get('style'):
                img['style'] = 'max-width: 100%; height: auto; display: block; margin: 1em auto;'
        
        # 处理表格
        for table in soup.find_all('table'):
            table['style'] = 'width: 100%; border-collapse: collapse; margin: 1em 0;'
        
        # 处理代码块
        for pre in soup.find_all('pre'):
            pre['style'] = 'background-color: #2c3e50; color: #ecf0f1; padding: 1em; border-radius: 5px; overflow-x: auto;'
        
        # 处理引用块
        for blockquote in soup.find_all('blockquote'):
            blockquote['style'] = 'margin: 1em 0; padding: 0.5em 1em; background-color: #f8f9fa; border-left: 4px solid #3498db;'
        
        optimized_html = str(soup)
        
        # 构建完整的HTML文档
        full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title if title else '微信公众号文章'}</title>
    {self.wechat_styles}
</head>
<body>
"""
        
        if title:
            full_html += f'    <h1 class="wechat-title">{title}</h1>\n'
        if subtitle:
            full_html += f'    <p class="wechat-subtitle">{subtitle}</p>\n'
        
        full_html += f"    {optimized_html}\n"
        full_html += """
</body>
</html>"""
        
        return full_html
    
    def convert_file(self, input_file, output_file=None, title="", subtitle=""):
        """转换文件"""
        try:
            # 检测文件格式
            file_format = self.detect_file_format(input_file)
            
            if file_format == 'unknown':
                print(f"不支持的文件格式: {Path(input_file).suffix}")
                return None
            
            print(f"📄 检测到文件格式: {file_format}")
            
            # 读取文件内容
            if file_format == 'docx':
                # Word文档需要特殊处理
                markdown_content = self.docx_to_markdown(input_file)
            else:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                markdown_content = self.convert_to_markdown(content, file_format, input_file)
            
            print(f"✅ 已转换为Markdown格式")
            
            # 转换为微信公众号HTML
            wechat_html = self.markdown_to_wechat_html(markdown_content, title, subtitle)
            
            # 确定输出文件名
            if not output_file:
                input_path = Path(input_file)
                output_file = input_path.with_suffix('.html')
            
            # 保存HTML文件
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(wechat_html)
            
            print(f"🎉 转换完成！")
            print(f"输入文件: {input_file} ({file_format})")
            print(f"输出文件: {output_file}")
            print(f"转换流程: {file_format} → Markdown → 微信公众号HTML")
            print(f"可以直接复制HTML内容到微信公众号编辑器")
            
            return wechat_html
            
        except Exception as e:
            print(f"❌ 转换失败: {str(e)}")
            return None


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='通用格式到微信公众号文章格式转换器')
    parser.add_argument('input', nargs='?', help='输入文件路径')
    parser.add_argument('-o', '--output', help='输出HTML文件路径')
    parser.add_argument('-t', '--title', help='文章标题')
    parser.add_argument('-s', '--subtitle', help='文章副标题')
    parser.add_argument('--style', help='文章风格', 
                       choices=WeChatStyleTemplates.get_available_styles(),
                       default='default')
    parser.add_argument('--list-styles', action='store_true', help='列出所有可用风格')
    parser.add_argument('--list-formats', action='store_true', help='列出支持的输入格式')
    
    args = parser.parse_args()
    
    # 列出支持的格式
    if args.list_formats:
        print("支持的输入格式：")
        print("=" * 50)
        formats = [
            ("Markdown", ".md, .markdown", "原生支持，功能最完整"),
            ("HTML", ".html, .htm", "转换为Markdown后处理"),
            ("纯文本", ".txt", "智能识别格式后转换"),
            ("RST", ".rst", "需要安装 docutils"),
            ("Word", ".docx", "需要安装 python-docx"),
            ("RTF", ".rtf", "需要安装 striprtf"),
        ]
        
        for name, extensions, description in formats:
            print(f"{name:10} {extensions:15} - {description}")
        
        print("\n转换流程：")
        print("其他格式 → Markdown → 微信公众号HTML")
        print("\n依赖安装命令：")
        print("pip install docutils python-docx striprtf")
        return
    
    # 列出风格
    if args.list_styles:
        print("可用的文章风格：")
        print("=" * 50)
        for style in WeChatStyleTemplates.get_available_styles():
            description = WeChatStyleTemplates.get_style_description(style)
            print(f"{style:12} - {description}")
        return
    
    # 检查输入文件
    if not args.input:
        parser.error("需要提供输入文件路径")
    
    # 创建转换器
    converter = UniversalToWeChatConverter(style=args.style)
    
    # 执行转换
    converter.convert_file(args.input, args.output, args.title, args.subtitle)


if __name__ == "__main__":
    main()
